import os
import re
from pathlib import Path
from enum import Enum

from tqdm import tqdm
from docx.api import Document as Doc
from bs4 import BeautifulSoup

from .util import normalize_spaces, is_bold, has_not_fewer_dots_than, drop_space_around_punctuation, is_h1, is_space, get_paragraph_style, get_document_comments
from .util.soup import get_first_non_empty_element
from .Table import Table
from .Paragraph import Paragraph
from .TableType import TableType
from .Document import Document
from .ContextRanker import ContextRanker, DEFAULT_MODEL as DEFAULT_EMBEDDING_MODEL


PARAGRAPH_SEP_PLACEHOLDER = '__PARAGRAPH_SEP__'
PARAGRAPH_SEP_DISTILLATION_PATTERN = re.compile(rf'\s*{PARAGRAPH_SEP_PLACEHOLDER}\s*')
PARAGRAPH_SEP = '\n\n'
TABLE_ID = re.compile('[A-ZА-Я0-9.Ё]+')
APPLICATION_ID = re.compile('[A-ZА-ЯЁ]+')
APPLICATION_TABLE_ID = re.compile('([A-ZА-ЯЁ]+).+')
NOT_APPLICATION_TABLE_ID = re.compile(r'\w+')
EXTERNAL_APPLICATION_REFERENCE_PATTERN = re.compile(r'.+сп\s+[0-9.]+\.?$')

# PARAGRAPH_ID = re.compile(r'\s*xmlns:[0-9a-z]*="[^"]+"')
PARAGRAPH_ID = re.compile(r'para[Ii]d="([0-9a-fA-F]+)"')
SEP = '<sep>'
PARAGRAPH = '{paragraph}'
TABLE = '{table}'

CRITERIA = (
    'Criteria for identifying paragraph, essential for understanding table content by decreasing importance:\n'
    'Direct Reference to the Table: Paragraph explicitly references the table by number, title, or content.\n'
    'Terminological Correlation: Paragraph shares key terms or synonyms with the table, indicating relevance through common terminology.\n'
    'Semantic Correlation: Paragraph content closely aligns semantically with the table content, addressing similar subjects or contexts even without explicit term overlap.\n'
    'Generalization of Table Content: Paragraph summarizes or generalizes all or part of the table content, providing an overview or conclusions derived from the data presented.\n'
    'Detailing Properties Relevant to Table Terms and Entities: Paragraph provides additional details about properties, characteristics, or attributes of key terms/entities'
    'mentioned in the table essential for understanding its data.\n'
    'Explanation of Abbreviations and Generic Terms: Paragraph clarifies abbreviations, acronyms, or general terms used within the table.\n'
    'Description of Table Structure and Layout:'
    'Paragraph describes the structure (rows, columns), organization, grouping of elements, headers, and relative positions of components within the table.\n'
    'Connections to Other Document Elements: Paragraph explains relationships between table content and other document sections (e.g., other tables, appendices, standards).\n'
    'Generalization or Summarization of Table Content: Paragraph summarizes or generalizes entire table content or significant portions thereof.\n'
    'Distance Between Table and Paragraph: Closer proximity (fewer intervening paragraphs/tables) typically indicates higher relevance.\n'
)

STEPS = (
    'Step-by-Step Procedure for applying these criteria to rank paragraphs:'
    'Step 1 Identify Direct References. First identify paragraphs explicitly referencing the table by number or title. These paragraphs usually have the highest priority.\n'
    'Step 2 Evaluate Terminological and Semantic Correlation.'
    'Identify paragraphs containing key terms identical or synonymous to those appearing in the table.'
    'Evaluate semantic similarity between paragraph content and table content; prioritize paragraphs clearly discussing similar topics.\n'
    'Step 3 Identify Structural and Explanatory Context. Select paragraphs describing:\n'
    '  - Table structure (rows, columns, sections).\n'
    '  - Definitions or explanations of abbreviations, generic terms used in the table.\n'
    '  - Detailed properties or characteristics of key concepts/entities listed in the table.\n'
    'Step 4 Determine Generalizations and Summaries. Identify paragraphs summarizing overall meaning or key insights derived from the entire table or substantial parts thereof.\n'
    'Step 5 Analyze Connections to Broader Document Context. Identify paragraphs connecting table content explicitly with:\n'
    '  - Other tables.\n'
    '  - Standards, regulations, appendices.\n'
    '  - External references mentioned in the document.\n'
    'Step 6 Consider Distance as Secondary Criterion. After applying previous steps, use paragraph proximity as a secondary criterion to fine-tune rankings:\n'
    '  - Closer paragraphs are generally more relevant when other criteria are similar in strength.\n'
)


def get_element_id(element):
    if (id_match := PARAGRAPH_ID.search(element)) is not None:
        id_ = id_match.group(1)
        return id_.lower()


def find_comment(element, comments):
    for el, comment in comments:
        if get_element_id(el) == get_element_id(str(element)):
            return comment


def join_paragraphs(paragraphs: list):
    if len(paragraphs) < 1:
        return None

    return PARAGRAPH_SEP_DISTILLATION_PATTERN.sub(
        PARAGRAPH_SEP_PLACEHOLDER,
        normalize_spaces(
            PARAGRAPH_SEP_PLACEHOLDER.join(paragraphs)
        )
    ).replace(PARAGRAPH_SEP_PLACEHOLDER, PARAGRAPH_SEP)


def extract_id(id_, pattern, text: str):
    for match in pattern.findall(text):
        id_candidate = str(match)

        if id_ is None or len(id_candidate) >= len(id_):
            id_ = id_candidate

    return id_


def is_not_part_of_other_id(text, id_, verbose = False):
    if verbose:
        print('checking for table id singularity')

    loc = text.find(id_)

    if loc > 0:
        prev_loc = loc - 1
        prev_char = None if prev_loc < 0 else text[prev_loc]

        next_loc = loc + len(id_)
        next_char = None if next_loc >= len(text) else text[next_loc]

        next_next_loc = next_loc + 1
        next_next_char = None if next_next_loc >= len(text) else text[next_next_loc]

        prefix_loc = prev_loc - 1
        prefix_char = None if prefix_loc < 1 else text[prefix_loc]
        prefix = None

        if prefix_char is not None:
            prefix = prefix_char

            while prefix_loc >= 0:
                prefix_loc -= 1
                prefix_char = None if prefix_loc < 0 else text[prefix_loc]

                if prefix_char is None or is_space(prefix_char):
                    break

                prefix += prefix_char

        prefix = None if prefix is None else prefix[::-1]

        if prefix is None:
            return False

        if verbose:
            # print(id_, '|', text[loc:])
            print(next_next_char)

        # print(text, '>>', prefix)
        normalized_prefix = prefix.lower().strip()

        # print(
        #     (
        #         normalized_prefix.startswith('прилож') or
        #         normalized_prefix.startswith('табл') or
        #         normalized_prefix.startswith('форм')
        #     ), prev_char is not None, (
        #         is_space(prev_char) or prev_char in '([{'
        #     ), next_char is not None, (
        #         is_space(next_char) or (
        #             next_char in ')]}' or
        #             next_char in '.' and next_next_char is None or
        #             next_next_char is not None and next_char in ';.,' and is_space(next_next_char)
        #         )
        #     )
        # )

        return (
            normalized_prefix.startswith('прилож') or
            normalized_prefix.startswith('табл') or
            normalized_prefix.startswith('форм')
        ) and prev_char is not None and (
            is_space(prev_char) or prev_char in '([{'
        ) and next_char is not None and (
            is_space(next_char) or (
                next_char in ')]}' or
                next_char in '.' and next_next_char is None or
                next_next_char is not None and next_char in ';.,' and is_space(next_next_char)
            )
        )


class Parser:
    def __init__(self, context_window_size: int = 5, json_indent: int = 2):
        self.context_window_size = context_window_size
        self.json_indent = json_indent

    def has_reference(self, text: str, table: Table, verbose: bool = False):
        id_ = table.id
        table_type = TableType(table.type)

        if id_ is None or NOT_APPLICATION_TABLE_ID.fullmatch(id_):
            application_table_id_match = None
        else:
            application_table_id_match = APPLICATION_TABLE_ID.fullmatch(id_)

        normalized_text = text.lower().strip()

        if verbose:
            print(is_not_part_of_other_id(text, id_, verbose = verbose))
            # print(
            #     id_ in text and is_not_part_of_other_id(text, id_, verbose = verbose),  # either there is a complete id in the text
            #     table_type == TableType.FORM,  # either the table looks like a form
            #     (
            #         application_table_id_match is not None and
            #         (
            #             re.search(r'\s' + application_table_id_match.group(1) + r'[^\w\s\.]', text) is not None
            #         ) and
            #         not text.endswith(application_table_id_match.group(1))
            #     )  # either there is an imcomplete reference (to the application which contains the table)
            # )

        return id_ is not None and not normalized_text.startswith('табл') and (  # text doesn't look like table description
            (
                table_type == TableType.TABLE and 'табл' in normalized_text or  # either table looks like a regular table and there is a stem 'tabl' in the given text
                table_type == TableType.FORM and ' форм' in normalized_text or  # either table looks like a form and there is a stem 'form' in the given text
                (
                    table_type == TableType.APPLICATION or
                    application_table_id_match is not None
                ) and (
                    'приложен' in normalized_text or 'табл' in normalized_text
                ) and (
                    EXTERNAL_APPLICATION_REFERENCE_PATTERN.fullmatch(normalized_text) is None
                )  # either table looks like a kind of application, and there is a stem 'applic' in the given text
            ) and
            (
                id_ in text and is_not_part_of_other_id(text, id_, verbose = verbose) or  # either there is a complete id in the text
                table_type == TableType.FORM or  # either the table looks like a form
                (
                    application_table_id_match is not None and
                    (
                        re.search(r'\s' + application_table_id_match.group(1) + r'[^\w\s\.]', text) is not None
                    ) and
                    not text.endswith(application_table_id_match.group(1))
                )  # either there is an imcomplete reference (to the application which contains the table)
            )
        )

        # return not normalized_text.startswith('табл') and (
        #     id_ is not None and
        #     (
        #         table_type == TableType.TABLE and (
        #             'табл' in normalized_text and there_are_paragraphs_with_full_id or
        #             'приложен' in normalized_text and not there_are_paragraphs_with_full_id
        #         ) or
        #         (
        #             table_type == TableType.APPLICATION or
        #             application_table_id_match is not None
        #         ) and (
        #             not there_are_paragraphs_with_full_id and 'приложен' in normalized_text or
        #             there_are_paragraphs_with_full_id and 'табл' in normalized_text
        #         ) and (
        #             EXTERNAL_APPLICATION_REFERENCE_PATTERN.fullmatch(normalized_text) is None
        #         ) or
        #         table_type == TableType.FORM and ' форм' in normalized_text
        #     ) and
        #     (
        #         id_ in text or
        #         table_type == TableType.FORM or
        #         (
        #             application_table_id_match is not None and
        #             (
        #                 not there_are_paragraphs_with_full_id and
        #                 re.search(r'\s' + application_table_id_match.group(1) + r'[^\w\s]', text) is not None
        #             ) and
        #             not text.endswith(application_table_id_match.group(1))
        #         )
        #     )
        # )

    def get_title(self, paragraphs, verbose: bool = False, title: str = None):
        # window = self.context_window_size
        # section_title = title

        title = []
        non_empty_paragraphs = []

        id_ = None
        table_type = None
        # full_reference_exists = False

        for j, paragraph in enumerate(paragraphs):
            text = '' if paragraph.text is None else drop_space_around_punctuation(normalize_spaces(paragraph.text))  # TODO: Add better handler for None values

            if len(text) > 0:
                non_empty_paragraphs.append(text)

            try:
                previous_paragraphs = paragraph.findPreviousSiblings('w:p')
            except AttributeError:
                previous_paragraphs = paragraphs[:j]

            try:
                next_paragraphs = paragraph.findNextSiblings('w:p')
            except AttributeError:
                next_paragraphs = paragraphs[j + 1:]

            if len(text.strip()) > 0:
                normalized_text = text.lower().strip()

                if normalized_text.startswith('библиография'):
                    title.append(text)
                    id_ = text
                    table_type = TableType.TABLE
                elif id_ is None and normalized_text.startswith('таблица'):
                    title.append(text)

                    for match in TABLE_ID.findall(text):
                        id_candidate = str(match)

                        if id_ is None or len(id_candidate) >= len(id_) and has_not_fewer_dots_than(id_candidate, id_):
                            id_ = id_candidate

                            try:
                                int(id_)
                                break
                            except ValueError:
                                pass

                    # there_are_paragraphs_with_full_id = any(id_ in paragraph.text for paragraph in previous_paragraphs)
                    application_table_id_match = APPLICATION_TABLE_ID.fullmatch(id_)

                    if application_table_id_match is not None and text.endswith(id_):
                        last_paragraph = get_first_non_empty_element(previous_paragraphs)
                        table_type = TableType.APPLICATION

                        if last_paragraph is not None and (
                            (
                                is_bold(paragraph) and (is_bold(last_paragraph) or is_h1(last_paragraph))
                            ) or (
                                not is_bold(paragraph) and not (is_bold(last_paragraph) or is_h1(last_paragraph))
                            )
                        ):
                            title.append(drop_space_around_punctuation(normalize_spaces(last_paragraph.text)))
                    else:
                        table_type = TableType.TABLE
                elif id_ is None and normalized_text.startswith('форма'):
                    title.append(text)

                    next_non_empty_paragraph = get_first_non_empty_element(next_paragraphs)
                    if next_non_empty_paragraph is not None:
                        title.append(next_non_empty_paragraph.text)

                    id_ = extract_id(id_, TABLE_ID, text)

                    table_type = TableType.FORM
                elif id_ is None and normalized_text.startswith('приложение'):
                    title.append(text)

                    next_non_empty_paragraph = get_first_non_empty_element(next_paragraphs)
                    if next_non_empty_paragraph is not None:
                        title.append(next_non_empty_paragraph.text)

                    id_ = extract_id(id_, APPLICATION_ID, text)

                    table_type = TableType.APPLICATION

        if 0 < len(non_empty_paragraphs) < 2 and len(title) < 1:
            # title = [section_title, *non_empty_paragraphs]
            title = non_empty_paragraphs
            table_type = TableType.TABLE

        return join_paragraphs(title), id_, TableType.TABLE if table_type is None else table_type

    def parse_file(self, source: str, get_destination: callable = None, cpu: bool = False, embedding_model: str = DEFAULT_EMBEDDING_MODEL):
        document = Doc(source)

        comments = get_document_comments(source)

        soup = BeautifulSoup(document._element.xml, 'lxml')

        if get_destination is None:
            stem = Path(source).stem

            def get_destination(i: int):
                return f'{stem}.{i:04d}'.replace('-', '_') + '.json'

        last_table_xml = None
        i = 0

        items = []

        contexts = {}
        tables = {}

        for i, element in list(enumerate(soup.find_all(re.compile(r'w:(tbl|p)')))):
            if element.name in ('w:p', 'w:tbl'):  # if element is a table
                comment = find_comment(element, comments)

                # print(comment)

                if element.name == 'w:tbl':
                    if (table := Table.from_soup(element, get_destination(i))):
                        items.append(table)
                        # print(table.json)

                        last_table_xml = str(element)

                        tables[comment] = table
                else:  # if element is a paragraph
                    if last_table_xml is not None:
                        if str(element) in last_table_xml:
                            continue
                        else:
                            last_table_xml = None

                    if (paragraph := Paragraph.from_soup(element)) is not None:
                        items.append(paragraph)

                    if comment is not None:
                        table_id, relevance_score = comment.split(' ')
                        relevance_score = float(relevance_score)

                        if (context := contexts.get(table_id)) is None:
                            contexts[table_id] = [(relevance_score, paragraph)]
                        else:
                            context.append((relevance_score, paragraph))

        return items

        # document = Document(items)
        # ranker = ContextRanker(model = embedding_model, cuda = not cpu)

        # for table in document.tables:
        #     ranker.rank(table, document.paragraphs)

        # print(document)

        # for i, table in list(enumerate(soup.find_all('w:tbl', 'w:p'))):
        #     yield Table.from_soup(
        #         table, get_destination(i)
        #     )

    def parse(self, source: str, destination: str, cpu: bool, embedding_model: str = DEFAULT_EMBEDDING_MODEL, paragraphs_filter: list[int] = None):
        if not os.path.isdir(destination):
            os.makedirs(destination)

        for source_file in os.listdir(source):
            if paragraphs_filter is None:
                print('Here is a document content. Each block of text contains either paragraph either table content. Adjacent blocks of texts are separated by an empty line.')
                print(f'Each paragraph starts with label "{PARAGRAPH}" followed by paragraph id and each table starts with label "{TABLE}", followed by table id. Cell contents within a table row is separated with "{SEP}", and rows are separated with a linebreak.')
                print('Your task is to rank document paragraphs by importance for understanding the table content. A paragraph is characterized by high rank (makes up context for a table), if it is essential for understanding table content.')
                print('Output only lists of paragraph ids for each table by decreasing the order of importance for unerstanding the table content. Your response must be a json object with format {"contexts": list[{"table": int, "paragraphs": list[int]}]}')
                print('Please, make sure that list of paragraphs for each table is complete, meaning that it contains all paragraph ids from the provided data, ordered by decreasing importance for understanding the table content')
                print()
                print(CRITERIA)
                print(STEPS)

            paragraph_id = 0
            table_id = 0

            relevant_paragraphs = None if paragraphs_filter is None else []

            for item in self.parse_file(
                source = os.path.join(source, source_file),
                get_destination = lambda i: os.path.join(destination, f'{Path(source_file).stem}.{i:04d}'.replace(' ', '_')) + '.json',
                cpu = cpu,
                embedding_model = embedding_model
            ):
                if isinstance(item, Paragraph):
                    paragraph_id = paragraph_id + 1
                    if paragraphs_filter is None or paragraph_id in paragraphs_filter:
                        if paragraphs_filter is None:
                            print(PARAGRAPH, f'{paragraph_id:03d}', item.content)
                            print()
                        else:
                            relevant_paragraphs.append((item.content, paragraphs_filter.index(paragraph_id)))
                if isinstance(item, Table) and paragraphs_filter is None:
                    print(TABLE, f'{(table_id := table_id + 1):03d}', '\n'.join([SEP.join(row) for row in item.content]))
                    print()

            if relevant_paragraphs is not None:
                for paragraph in sorted(relevant_paragraphs, key = lambda item: item[1]):
                    print(paragraph[0])
